import os
import io
import requests
import sys

#--- It's time to start creating docx document 
#-- create docx-file
from docx import Document
from docx.shared import Inches 

s_category=sys.argv[3]

#-- set document filename
s_output_filename =  "News"+".docx"
s_docx_path=sys.argv[2]
s_docx_absl_path=os.path.abspath(s_docx_path)

exists = os.path.isfile(s_docx_absl_path+"/"+s_output_filename)
if exists:
	f = open(s_docx_absl_path+"/"+s_output_filename, 'rb')
	document = Document(f)
else:
	document = Document()


document.add_heading(s_category)

s_article_path=sys.argv[1]
s_art_absl_path=os.path.abspath(s_article_path)

sorted_dir = sorted(os.listdir(s_art_absl_path))

import newspaper
for s_file_name in os.listdir(s_art_absl_path):	
	s_absolute_path = os.path.abspath(s_art_absl_path+'/'+s_file_name)	
	print("Start parsing "+s_file_name+":")
	with open(s_absolute_path) as fp:
		html = fp.read()
		article = newspaper.Article(url='http://example.com/test-url')
		article.set_html(html)				
		article.parse()
		article.nlp()		
		
		s_article_title = article.title					
		
		response = requests.get(article.top_image, stream=True)
		b_image = io.BytesIO(response.content)		
		
		s_article_text = article.text
						
		
		# add article-title and (if available) -subtitle to docx
		document.add_heading(s_article_title, level=2)
		#document.add_heading(s_article_sub_title, level=2)
		
		document.add_picture(b_image,width=Inches(3.25))

		# add text content (paragraph) to docx
		paragraph = document.add_paragraph(s_article_text)

		document.add_page_break()
#-------------------------------------------------------------------------------------


# save it
print("Save document")
document.save(s_docx_absl_path+"/"+s_output_filename)
